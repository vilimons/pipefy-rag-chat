from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from fastapi import UploadFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError

SUPPORTED_FILE_EXTENSIONS = {".pdf", ".txt", ".docx"}


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


async def extract_text_from_upload(file: UploadFile) -> str:
    extension = _get_file_extension(file.filename)

    if extension not in SUPPORTED_FILE_EXTENSIONS:
        raise UnsupportedFileTypeError(
            "Tipo de arquivo não suportado. Envie apenas arquivos PDF, TXT ou DOCX."
        )

    content = await file.read()

    try:
        if extension == ".txt":
            text = _extract_text_from_txt(content)
        elif extension == ".pdf":
            text = _extract_text_from_pdf(content)
        else:
            text = _extract_text_from_docx(content)
    except (BadZipFile, PackageNotFoundError, PdfReadError) as error:
        raise EmptyDocumentError(
            "Não foi possível ler o documento. Verifique se o arquivo não "
            "está corrompido e foi salvo como PDF, TXT ou DOCX válido."
        ) from error

    normalized_text = normalize_text(text)

    if not normalized_text:
        raise EmptyDocumentError("O documento não contém texto legível.")

    return normalized_text


def _get_file_extension(filename: str | None) -> str:
    if not filename:
        return ""

    return Path(filename).suffix.lower()


def _extract_text_from_txt(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore")


def _extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))

    pages = [page.extract_text() or "" for page in reader.pages]

    return "\n".join(pages)


def _extract_text_from_docx(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    table_cells: list[str] = []

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text)

    return "\n".join([*paragraphs, *table_cells])


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    return "\n".join(lines)
